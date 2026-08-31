# -*- coding: utf-8 -*-
"""Korporativ mijoz uchun oylik hujjatlar (Word).

Har oy oxirida buxgalteriya ikki hujjat so'raydi:

  * **Bajarilgan ishlar dalolatnomasi** — davr ichida qancha xizmat
    ko'rsatilgani. U hisob-fakturaning asosi bo'ladi.
  * **Solishtirma dalolatnoma** — davr boshidagi va oxiridagi qoldiq,
    o'rtadagi kirim-chiqim. Ikki tomonning hisobi bir xilligini
    tasdiqlaydi.

Nima uchun generatsiya, qo'lda yozish emas: bir oyda yuzlab sessiya
bo'ladi va ularni qo'lda yig'ish xatoga olib keladi. Xato esa nizoga
aylanadi va odatda biz zarar ko'ramiz.

Ma'lumot manbai — sessiyalar va hamyon tranzaksiyalari. Ular kompaniyaning
hisob foydalanuvchisiga (`billing_user`) bog'langan.
"""

from datetime import date, timedelta
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from dashboard.banking import format_account, format_inn
from dashboard.invoices import amount_in_words
from dashboard.templatetags.money import format_som

MONTHS = ['', 'yanvar', 'fevral', 'mart', 'aprel', 'may', 'iyun',
          'iyul', 'avgust', 'sentabr', 'oktabr', 'noyabr', 'dekabr']


def month_range(year: int, month: int):
    """Oyning birinchi va oxirgi kuni."""
    start = date(year, month, 1)
    end = date(year + (month == 12), month % 12 + 1, 1) - timedelta(days=1)
    return start, end


def month_label(year: int, month: int) -> str:
    return f'{year}-yil {MONTHS[month]}'


class _Builder:
    """Ikkala hujjat uchun umumiy qism: uslub, sarlavha, imzo."""

    def __init__(self, company, settings_obj, start, end):
        self.company = company
        self.settings = settings_obj
        self.start = start
        self.end = end
        self.doc = Document()
        self._setup()

    def _setup(self):
        style = self.doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(10)
        for section in self.doc.sections:
            section.left_margin = Cm(2)
            section.right_margin = Cm(1.5)
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)

    def _para(self, text, bold=False, align=None, size=None):
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        if size:
            run.font.size = Pt(size)
        if align is not None:
            para.alignment = align
        para.paragraph_format.space_after = Pt(2)
        return para

    @property
    def supplier(self):
        return self.settings.org_legal_name or self.settings.app_name

    def title_block(self, title):
        self._para(title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13)
        self._para(f'{self.start:%d.%m.%Y} — {self.end:%d.%m.%Y} davri uchun',
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        self.doc.add_paragraph()

    def parties_block(self):
        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        rows = [
            ('Ijrochi', self.supplier),
            ('STIR', format_inn(self.settings.org_inn)),
            ('H/r', format_account(self.settings.org_bank_account)),
            ('Buyurtmachi', self.company.invoice_name),
            ("Buyurtmachi STIR", format_inn(self.company.inn)),
            ("Buyurtmachi h/r", format_account(self.company.bank_account)),
        ]
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].paragraphs[0].add_run(label).bold = True
            cells[1].text = (value or '').strip() or '_' * 18
        self.doc.add_paragraph()

    def signature_block(self):
        director = (self.settings.org_director or '').strip() or '_' * 18
        client = (self.company.director or '').strip() or '_' * 18

        self.doc.add_paragraph()
        table = self.doc.add_table(rows=1, cols=2)
        left, right = table.rows[0].cells
        left.paragraphs[0].add_run('IJROCHI').bold = True
        left.add_paragraph(self.supplier)
        left.add_paragraph()
        left.add_paragraph(f'__________________ {director}')
        left.add_paragraph("M.O'.")

        right.paragraphs[0].add_run('BUYURTMACHI').bold = True
        right.add_paragraph(self.company.invoice_name)
        right.add_paragraph()
        right.add_paragraph(f'__________________ {client}')
        right.add_paragraph("M.O'.")

    def save(self) -> BytesIO:
        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer


class ActBuilder(_Builder):
    """Bajarilgan ishlar dalolatnomasi."""

    def __init__(self, company, settings_obj, start, end, sessions):
        super().__init__(company, settings_obj, start, end)
        self.sessions = list(sessions)

    def services_block(self):
        energy = sum(s.kwh_charged or 0 for s in self.sessions)
        energy_cost = sum(s.energy_cost or 0 for s in self.sessions)
        parking_cost = sum(s.parking_cost or 0 for s in self.sessions)
        total = energy_cost + parking_cost

        table = self.doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        for cell, title in zip(table.rows[0].cells,
                               ['№', 'Xizmat nomi', "O'lchov", 'Miqdor', "Summa (so'm)"]):
            cell.paragraphs[0].add_run(title).bold = True

        row = table.add_row().cells
        row[0].text = '1'
        row[1].text = 'Elektromobillarni zaryadlash xizmati'
        row[2].text = 'kVt·soat'
        row[3].text = f'{energy:.2f}'
        row[4].text = format_som(energy_cost)

        # Parkovka faqat bo'lgan taqdirda ko'rsatiladi: nol qatori
        # hujjatni uzaytiradi va savol tug'diradi
        if parking_cost:
            row = table.add_row().cells
            row[0].text = '2'
            row[1].text = 'Ulagichni band qilib turish (parkovka)'
            row[2].text = 'daqiqa'
            row[3].text = str(sum(s.parking_minutes or 0 for s in self.sessions))
            row[4].text = format_som(parking_cost)

        total_row = table.add_row().cells
        total_row[3].paragraphs[0].add_run('Jami:').bold = True
        total_row[4].paragraphs[0].add_run(format_som(total)).bold = True

        self.doc.add_paragraph()
        self._para(f"Jami: {format_som(total)} so'm ({amount_in_words(total)} so'm).",
                   bold=True)
        self._para(f'Davr ichida {len(self.sessions)} ta zaryadlash sessiyasi '
                   f"bo'lib o'tdi.", size=9)
        self.doc.add_paragraph()
        self._para('Tomonlar bir-biriga da\'vo qilmaydi. Xizmat to\'liq va sifatli '
                   "ko'rsatildi.", size=9)
        return total

    def build(self):
        self.title_block('BAJARILGAN ISHLAR DALOLATNOMASI')
        self.parties_block()
        self.services_block()
        self.signature_block()
        return self.save()


class ReconciliationBuilder(_Builder):
    """Solishtirma dalolatnoma — davr boshidagi va oxiridagi qoldiq."""

    def __init__(self, company, settings_obj, start, end, rows, opening, closing):
        super().__init__(company, settings_obj, start, end)
        self.rows = list(rows)
        self.opening = opening
        self.closing = closing

    def movements_block(self):
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        for cell, title in zip(table.rows[0].cells,
                               ['Sana', 'Asos', 'Kirim', 'Chiqim']):
            cell.paragraphs[0].add_run(title).bold = True

        opening = table.add_row().cells
        opening[1].paragraphs[0].add_run('Davr boshiga qoldiq').bold = True
        opening[2].text = format_som(self.opening)

        debit = credit = 0
        for record in self.rows:
            cells = table.add_row().cells
            cells[0].text = f'{record.created_at:%d.%m.%Y}'
            cells[1].text = record.description or record.get_type_display()
            if record.type == record.Type.TOPUP:
                cells[2].text = format_som(record.amount)
                debit += record.amount
            else:
                cells[3].text = format_som(record.amount)
                credit += record.amount

        totals = table.add_row().cells
        totals[1].paragraphs[0].add_run('Davr aylanmasi').bold = True
        totals[2].paragraphs[0].add_run(format_som(debit)).bold = True
        totals[3].paragraphs[0].add_run(format_som(credit)).bold = True

        closing = table.add_row().cells
        closing[1].paragraphs[0].add_run('Davr oxiriga qoldiq').bold = True
        closing[2].paragraphs[0].add_run(format_som(self.closing)).bold = True

        self.doc.add_paragraph()
        self._para(
            f"Davr oxiriga Buyurtmachining oldindan to'langan qoldig'i "
            f"{format_som(self.closing)} so'm ({amount_in_words(self.closing)} so'm).",
            bold=True)
        self._para('Tomonlar hisob-kitobni solishtirdilar va yuqoridagi '
                   "ma'lumotlar bilan roziligini tasdiqlaydilar.", size=9)

    def build(self):
        self.title_block('SOLISHTIRMA DALOLATNOMA')
        self.parties_block()
        self.movements_block()
        self.signature_block()
        return self.save()


# ── Ma'lumot yig'ish ────────────────────────────────────────────
def period_sessions(company, start, end):
    """Davr ichida TUGAGAN sessiyalar.

    Ketayotgan sessiya hujjatga kirmaydi: uning yakuniy summasi hali
    ma'lum emas va u keyingi davrga tushadi.
    """
    from sessions_app.models import ChargingSession

    return (ChargingSession.objects
            .filter(user=company.billing_user,
                    stopped_at__date__gte=start, stopped_at__date__lte=end)
            .exclude(status=ChargingSession.Status.CHARGING)
            .select_related('station')
            .order_by('stopped_at'))


def period_movements(company, start, end):
    """Davr ichidagi hamyon harakatlari."""
    from wallet.models import Transaction

    return (Transaction.objects
            .filter(user=company.billing_user,
                    created_at__date__gte=start, created_at__date__lte=end)
            .order_by('created_at'))


def balance_at(company, moment_date):
    """Berilgan kun BOSHIDAGI qoldiq.

    Hozirgi balansdan keyingi harakatlarni ayirib hisoblanadi: hamyonda
    faqat joriy qoldiq saqlanadi, tarix esa tranzaksiyalarda.
    """
    from wallet.models import Transaction

    balance = company.balance
    later = Transaction.objects.filter(
        user=company.billing_user, created_at__date__gte=moment_date)
    for record in later:
        if record.type == record.Type.TOPUP:
            balance -= record.amount
        else:
            balance += record.amount
    return max(0, balance)


def build_act(company, year, month):
    from management.models import SiteSettings

    start, end = month_range(year, month)
    return ActBuilder(company, SiteSettings.load(), start, end,
                      period_sessions(company, start, end)).build()


def build_reconciliation(company, year, month):
    from management.models import SiteSettings

    start, end = month_range(year, month)
    rows = period_movements(company, start, end)
    opening = balance_at(company, start)

    closing = opening
    for record in rows:
        if record.type == record.Type.TOPUP:
            closing += record.amount
        else:
            closing -= record.amount

    return ReconciliationBuilder(company, SiteSettings.load(), start, end,
                                 rows, opening, max(0, closing)).build()
