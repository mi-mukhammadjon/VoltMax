# -*- coding: utf-8 -*-
"""To'lov uchun hisob (schyot-faktura) — Word hujjati.

Korporativ mijoz bank orqali to'laydi, buning uchun buxgalteriyasiga rasmiy
hisob kerak: kimga, qaysi hisob raqamiga, qancha va nima uchun. Hujjat har
safar qaytadan yig'iladi — rekvizitlar sozlamalardan va mijoz kartochkasidan
olinadi, shuning uchun ular o'zgarsa hisob ham o'zi yangilanadi.
"""

from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from dashboard.banking import format_account, format_inn
from dashboard.phones import format_phone
from dashboard.templatetags.money import format_som

# ── Summani so'z bilan yozish ───────────────────────────────────
# Hisobda "summa so'z bilan" majburiy: raqamga bitta nol qo'shib qo'yish
# oson, so'z bilan yozilgani esa buni darrov ko'rsatadi.
_ONES = ['', 'bir', 'ikki', 'uch', "to'rt", 'besh', 'olti', 'yetti', 'sakkiz', "to'qqiz"]
_TENS = ['', "o'n", 'yigirma', "o'ttiz", 'qirq', 'ellik', 'oltmish', 'yetmish',
         'sakson', "to'qson"]
_SCALES = [(10 ** 9, 'milliard'), (10 ** 6, 'million'), (1000, 'ming')]


def _under_thousand(value):
    words = []
    if value >= 100:
        words += [_ONES[value // 100], 'yuz']
        value %= 100
    if value >= 10:
        words.append(_TENS[value // 10])
        value %= 10
    if value:
        words.append(_ONES[value])
    return words


def amount_in_words(amount: int) -> str:
    """123456 → «bir yuz yigirma uch ming to'rt yuz ellik olti»."""
    amount = int(amount)
    if amount == 0:
        return 'nol'

    words = []
    for scale, name in _SCALES:
        if amount >= scale:
            chunk = amount // scale
            # «bir ming» emas, shunchaki «ming» deyiladi
            if not (chunk == 1 and scale == 1000):
                words += _under_thousand(chunk)
            words.append(name)
            amount %= scale
    words += _under_thousand(amount)
    return ' '.join(word for word in words if word)


class InvoiceBuilder:
    """Hisob hujjatini yig'adi."""

    def __init__(self, invoice, settings_obj):
        self.invoice = invoice
        self.company = invoice.company
        self.settings = settings_obj
        self.doc = Document()
        self._setup_styles()

    def _setup_styles(self):
        style = self.doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(10)
        for section in self.doc.sections:
            section.left_margin = Cm(2)
            section.right_margin = Cm(1.5)
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)

    def _para(self, text, bold=False, align=None, size=None):
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        if size:
            run.font.size = Pt(size)
        if align is not None:
            para.alignment = align
        para.paragraph_format.space_after = Pt(2)
        return para

    # ── Bo'limlar ───────────────────────────────────────────────
    def header_block(self):
        invoice = self.invoice
        self._para(
            f"To'lov uchun hisob № {invoice.number} "
            f"({invoice.issued_at:%d.%m.%Y})",
            bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=13,
        )
        self.doc.add_paragraph()
        self._para(
            "Diqqat! Tovar (xizmat) faqat to'lov amalga oshirilgandan so'ng "
            "beriladi. Ushbu hisob bo'yicha to'lov qilinishi shartnoma "
            "shartlariga rozilik deb hisoblanadi.",
            size=9,
        )
        self.doc.add_paragraph()

    def parties_block(self):
        settings_obj = self.settings
        company = self.company

        table = self.doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'

        rows = [
            ("Yetkazib beruvchi", settings_obj.org_legal_name or settings_obj.app_name),
            ('Manzil', settings_obj.org_address),
            ('STIR', format_inn(settings_obj.org_inn)),
            ('Bank', settings_obj.org_bank_name),
            ("Hisob raqami", format_account(settings_obj.org_bank_account)),
            ('MFO', settings_obj.org_bank_mfo),
            ('Telefon', format_phone(settings_obj.support_phone)),
            ("To'lovchi", company.invoice_name),
            ("To'lovchi manzili", company.legal_address),
            ("To'lovchi STIR", format_inn(company.inn)),
            ("To'lovchi banki", company.bank_name),
            ("To'lovchi h/r", format_account(company.bank_account)),
            ("To'lovchi MFO", company.bank_mfo),
        ]
        for label, value in rows:
            cells = table.add_row().cells
            cells[0].paragraphs[0].add_run(label).bold = True
            cells[1].text = (value or '').strip() or '_' * 20
        self.doc.add_paragraph()

    def items_block(self):
        invoice = self.invoice
        table = self.doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        headers = ['№', 'Xizmat nomi', "O'lchov", 'Miqdor', "Summa (so'm)"]
        for cell, title in zip(table.rows[0].cells, headers):
            cell.paragraphs[0].add_run(title).bold = True

        row = table.add_row().cells
        row[0].text = '1'
        row[1].text = invoice.purpose
        row[2].text = 'xizmat'
        row[3].text = '1'
        row[4].text = format_som(invoice.amount)

        total = table.add_row().cells
        total[3].paragraphs[0].add_run('Jami:').bold = True
        total[4].paragraphs[0].add_run(format_som(invoice.amount)).bold = True

        self.doc.add_paragraph()
        self._para(
            f"To'lov summasi: {format_som(invoice.amount)} so'm "
            f"({amount_in_words(invoice.amount)} so'm).",
            bold=True,
        )
        # QQS bo'yicha holat rekvizitdan aniqlanadi: kod bo'lmasa soliq
        # ajratilmaydi va buni hujjatda ochiq yozib qo'yish kerak
        if not getattr(self.settings, 'org_vat_code', ''):
            self._para("QQS solinmaydi.", size=9)
        self.doc.add_paragraph()

    def signature_block(self):
        director = (self.settings.org_director or '').strip() or '_' * 20
        table = self.doc.add_table(rows=1, cols=2)
        left, right = table.rows[0].cells
        left.text = f'Rahbar ______________ {director}'
        right.text = 'Bosh hisobchi ______________'
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self.doc.add_paragraph()
        self._para("M.O'.", size=9)

    def build(self) -> BytesIO:
        self.header_block()
        self.parties_block()
        self.items_block()
        self.signature_block()

        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer


def build_invoice_document(invoice) -> BytesIO:
    """Hisob uchun `.docx` faylni qaytaradi."""
    from management.models import SiteSettings

    return InvoiceBuilder(invoice, SiteSettings.load()).build()
