# -*- coding: utf-8 -*-
"""Korporativ mijoz bilan xizmat ko'rsatish shartnomasi (Word).

Nima uchun generatsiya, tayyor fayl emas: shartnomada mijozning rekvizitlari,
tariflar va kartalar ro'yxati bo'ladi — ular bazada turadi va o'zgarib
boradi. Tayyor `.docx` faylni qo'lda to'ldirish esa xatoga olib keladi
(STIR yoki hisob raqamida bitta xato butun to'lovni to'xtatadi).

Hujjat SHABLON: bo'sh qolgan rekvizitlar `______` bo'lib chiqadi va
operator ularni qo'lda to'ldiradi. Shartnoma raqami va sanasi ham shunday —
ular buxgalteriya jurnalidan olinadi.
"""

import re
from io import BytesIO

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt

from dashboard.banking import format_account, format_inn
from dashboard.phones import format_phone
from dashboard.templatetags.money import format_som

BLANK = '_' * 18

def _text(value, width=18):
    """Bo'sh qiymat o'rniga to'ldirish uchun chiziq."""
    value = (value or '').strip()
    return value or '_' * width


class ContractBuilder:
    """Shartnoma hujjatini yig'adi.

    Har bo'lim alohida metod — matnni o'zgartirish kerak bo'lganda qaysi
    joyni tahrirlash kerakligi darrov ko'rinadi.
    """

    def __init__(self, company, settings_obj, cards, sections):
        self.company = company
        self.settings = settings_obj
        self.cards = list(cards)
        self.sections = list(sections)
        self.doc = Document()
        self._setup_styles()
        self.values = self._placeholder_values()

    # ── O'rin egallovchilar ─────────────────────────────────────
    def _placeholder_values(self):
        return {
            'ijrochi': _text(self.settings.org_legal_name or self.settings.app_name, 30),
            'ijrochi_rahbari': _text(self.settings.org_director, 24),
            'buyurtmachi': _text(self.company.invoice_name, 30),
            'buyurtmachi_rahbari': _text(self.company.director, 24),
            'narx': format_som(self.settings.default_price_per_kwh),
            'parkovka': format_som(self.settings.default_parking_fee_per_min),
            'kartalar_soni': str(len(self.cards)),
            'shahar': self.settings.contract_city,
        }

    def _fill(self, text):
        """`{narx}` kabi o'rin egallovchilarni qiymatga almashtiradi.

        Noma'lum nom o'z holicha qoladi — operator xato yozganini hujjatda
        ko'radi va tuzatadi. `format()` bo'lsa butun generatsiya qulab tushardi.
        """
        return re.sub(
            r'\{(\w+)\}',
            lambda match: self.values.get(match.group(1), match.group(0)),
            text or '',
        )

    # ── Umumiy ko'rinish ────────────────────────────────────────
    def _setup_styles(self):
        style = self.doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)

        for section in self.doc.sections:
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(1.5)
            section.top_margin = Cm(2)
            section.bottom_margin = Cm(2)

    def _heading(self, text):
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.bold = True
        para.paragraph_format.space_before = Pt(10)
        para.paragraph_format.space_after = Pt(4)
        return para

    def _para(self, text, bold=False, align=None):
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        if align is not None:
            para.alignment = align
        para.paragraph_format.space_after = Pt(3)
        return para

    def _item(self, number, text):
        para = self.doc.add_paragraph()
        para.add_run(f'{number}. ').bold = True
        para.add_run(text)
        para.paragraph_format.space_after = Pt(3)
        return para

    # ── Bo'limlar ───────────────────────────────────────────────
    def title_block(self):
        self._para(self.settings.contract_title,
                   bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        head = self.doc.add_paragraph()
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        head.add_run(f'№ {BLANK}')

        row = self.doc.add_table(rows=1, cols=2)
        row.autofit = True
        left, right = row.rows[0].cells
        left.text = self.settings.contract_city
        right.paragraphs[0].add_run(f'«{"__"}» {"_" * 12} 20{"__"} y.')
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self.doc.add_paragraph()

    def parties_block(self):
        """Kirish qismi — matni sozlamalardan olinadi, tomonlar nomi qalin."""
        para = self.doc.add_paragraph()
        names = [self.values['ijrochi'], self.values['buyurtmachi']]
        for chunk in self._split_bold(self._fill(self.settings.contract_preamble), names):
            run = para.add_run(chunk)
            run.bold = chunk in names
        para.paragraph_format.space_after = Pt(8)

    @staticmethod
    def _split_bold(text, names):
        """Matnni tomonlar nomi bo'yicha bo'laklarga ajratadi.

        Nom o'zgaruvchan bo'lgani uchun uni matn ichidan qidiramiz — operator
        preambulani qanday yozgan bo'lsa ham nomlar qalin bo'lib chiqadi.
        """
        parts = [text]
        for name in names:
            if not name:
                continue
            expanded = []
            for part in parts:
                if part in names:
                    expanded.append(part)
                    continue
                pieces = part.split(name)
                for index, piece in enumerate(pieces):
                    if index:
                        expanded.append(name)
                    if piece:
                        expanded.append(piece)
            parts = expanded
        return parts

    def sections_block(self):
        """Shartlar — bazadagi tahrirlanadigan bo'limlardan yig'iladi.

        Raqamlash saqlanmaydi, shu yerda hisoblanadi: bo'lim o'chirilsa yoki
        o'rni almashtirilsa qolganlari o'zi qayta raqamlanadi.
        """
        for number, section in enumerate(self.sections, 1):
            self._heading(f'{number}. {section.title.upper()}')

            item = 0       # joriy band raqami (N.M)
            sub = 0        # joriy ichki band raqami (N.M.K)
            for is_sub, text in section.items():
                text = self._fill(text)
                if is_sub:
                    # Ichki band oldingi bandsiz bo'lishi mumkin emas —
                    # bunday matnda uni oddiy band sifatida chiqaramiz
                    if item == 0:
                        item += 1
                        self._item(f'{number}.{item}', text)
                        continue
                    sub += 1
                    self._item(f'{number}.{item}.{sub}', text)
                else:
                    item += 1
                    sub = 0
                    self._item(f'{number}.{item}', text)
        return len(self.sections)

    def requisites_block(self, number):
        self._heading(f'{number}. TOMONLARNING REKVIZITLARI VA IMZOLARI')

        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        left, right = table.rows[0].cells
        self._fill_party(left, 'IJROCHI', {
            'Nomi': self.settings.org_legal_name or self.settings.app_name,
            'Manzil': self.settings.org_address,
            'STIR': format_inn(self.settings.org_inn),
            'Bank': self.settings.org_bank_name,
            'H/r': format_account(self.settings.org_bank_account),
            'MFO': self.settings.org_bank_mfo,
            'Tel': format_phone(self.settings.support_phone),
        }, self.settings.org_director)

        self._fill_party(right, 'BUYURTMACHI', {
            'Nomi': self.company.invoice_name,
            'Manzil': self.company.legal_address,
            'STIR': format_inn(self.company.inn),
            'Bank': self.company.bank_name,
            'H/r': format_account(self.company.bank_account),
            'MFO': self.company.bank_mfo,
            'Tel': format_phone(self.company.contact_phone),
        }, self.company.director)

    def _fill_party(self, cell, title, fields, director):
        cell.paragraphs[0].add_run(title).bold = True
        for label, value in fields.items():
            para = cell.add_paragraph()
            para.add_run(f'{label}: ')
            para.add_run(_text(value))
            para.paragraph_format.space_after = Pt(1)

        cell.add_paragraph()
        sign = cell.add_paragraph()
        sign.add_run('__________________ ')
        sign.add_run(_text(director, 20))
        cell.add_paragraph('M.O\'.')

    def cards_appendix(self):
        """1-ilova: mijozga biriktirilgan kartalar ro'yxati."""
        self.doc.add_page_break()
        self._para('1-ILOVA', bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
        self._para(f'№ {BLANK} sonli shartnomaga', align=WD_ALIGN_PARAGRAPH.RIGHT)
        self.doc.add_paragraph()
        self._para('BERILGAN RFID KARTALAR RO\'YXATI',
                   bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        self.doc.add_paragraph()

        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        for cell, title in zip(table.rows[0].cells,
                               ['№', 'Karta raqami', 'Nomi / mas\'ul shaxs', 'Holati']):
            cell.paragraphs[0].add_run(title).bold = True

        if self.cards:
            for index, card in enumerate(self.cards, 1):
                row = table.add_row().cells
                row[0].text = str(index)
                row[1].text = card.id_tag
                row[2].text = card.label or (card.user.username if card.user_id else '')
                row[3].text = card.effective_status_display
        else:
            # Kartalar hali biriktirilmagan — bo'sh qatorlar qoldiramiz,
            # operator ularni qo'lda to'ldiradi
            for index in range(1, 6):
                row = table.add_row().cells
                row[0].text = str(index)

        self.doc.add_paragraph()
        self._para(self._fill(self.settings.contract_appendix_note))
        self.doc.add_paragraph()

        sign = self.doc.add_table(rows=1, cols=2)
        left, right = sign.rows[0].cells
        left.text = 'Ijrochi: __________________'
        right.text = 'Buyurtmachi: __________________'
        right.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # ── Yig'ish ─────────────────────────────────────────────────
    def build(self) -> BytesIO:
        self.title_block()
        self.parties_block()
        last = self.sections_block()
        self.requisites_block(last + 1)
        self.cards_appendix()

        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer


def build_company_contract(company):
    """Kompaniya uchun shartnoma shablonini `.docx` sifatida qaytaradi.

    Bo'limlar bazadan olinadi — operator ularni «Sozlamalar > Shartnoma»
    bo'limida tahrirlaydi. Baza bo'sh bo'lsa standart matn yaratiladi,
    shunda hujjat hech qachon shartlarsiz chiqmaydi.
    """
    from management.models import ContractSection, SiteSettings

    ContractSection.ensure_defaults()
    cards = company.cards.select_related('user').order_by('id_tag')
    sections = ContractSection.objects.filter(is_active=True)
    return ContractBuilder(company, SiteSettings.load(), cards, sections).build()
